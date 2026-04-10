""" 
Forex Geopolitical Signal Intelligence App - Streamlit Frontend
=================================================================
Research agent: OpenAI gpt-4.1 via Responses API + web_search_preview
Economic data : FRED API (free, structured) +OpenAI web search (combined)
Synthesis agent: Anthropic Claude (claude-sonnet-4-6) -> FX intelligence briefing

Required .env variables:
- ANTHROPIC_API_KEY=your_anthropic_api_key
- OPENAI_API_KEY=your_openai_api_key
- FRED_API_KEY=your_fred_api_key    <- sign up for free at https://fred.stlouisfed.org/docs/api/api_key.html

FX Pairs covered: USD/GBP, USD/EUR, GBP/EUR
"""

import os
import json
import requests
from datetime import datetime, timedelta
from dotenv import load_dotenv
from pathlib import Path
import io

#load_dotenv(Path(__file__).parent / ".env")
load_dotenv(Path(__file__).parent / ".env", override=True, encoding='utf-8')
try:
    import streamlit as _st
    for key in ("ANTHROPIC_API_KEY", "OPENAI_API_KEY", "FRED_API_KEY"):
        if key in _st.secrets:
            os.environ[key] = _st.secrets[key]
except Exception:
    pass

import anthropic
import streamlit as st
from openai import OpenAI
from docx import Document
from docx.shared import Pt, RGBColor

# ------------------------------------------------------
# Page Configuration
# ------------------------------------------------------   
st.set_page_config(
    page_title="Forex Geopolitical Signal Intelligence",
    page_icon="O",
    layout="wide",
    initial_sidebar_state="expanded"
) 

#Load API keys form Streamlit secrets
for key in ["ANTHROPIC_API_KEY", "OPENAI_API_KEY", "FRED_API_KEY"]:
    if key in st.secrets:
        os.environ[key] = st.secrets[key]

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=IBM+Plex+Mono:wght@400;600&family=IBM+Plex+Sans:wght@300;400;600&display=swap');
 
    html, body, .stApp { background-color: #080b12; color: #d4dbe8; font-family: 'IBM Plex Sans', sans-serif; }
    .block-container { padding-top: 1.8rem; max-width: 1280px; }
 
    h1 { font-family: 'IBM Plex Mono', monospace; font-size: 1.4rem !important;
         color: #7eb8f7; letter-spacing: 0.05em; }
    h3 { font-size: 1rem !important; color: #a8bdd4; margin-bottom: 0.25rem; }
 
    .pair-card {
        background: #0d1825; border: 1px solid #1e3050; border-radius: 6px;
        padding: 14px 16px; text-align: center;
    }
    .pair-name  { color: #7eb8f7; font-family: 'IBM Plex Mono', monospace;
                  font-size: 0.85rem; letter-spacing: 0.08em; margin-bottom: 4px; }
    .pair-rate  { color: #d4dbe8; font-size: 1.5rem; font-weight: 600; }
    .pair-delta-up { color: #3cb878; font-size: 0.75rem; }
    .pair-delta-dn { color: #e05c5c; font-size: 0.75rem; }
 
    .econ-card {
        background: #060d18; border: 1px solid #142030; border-radius: 6px;
        padding: 10px 14px; margin-bottom: 6px;
    }
    .econ-label { color: #4a6a8a; font-size: 0.68rem; font-family: 'IBM Plex Mono', monospace;
                  letter-spacing: 0.07em; margin-bottom: 2px; }
    .econ-values { display: flex; gap: 18px; align-items: baseline; }
    .econ-country { color: #5a8ab5; font-size: 0.68rem; font-family: 'IBM Plex Mono', monospace; }
    .econ-val { color: #d4dbe8; font-size: 1rem; font-weight: 600; }
    .econ-prev { color: #4a6a8a; font-size: 0.72rem; margin-left: 4px; }
 
    .log-box {
        background: #060910; border: 1px solid #141c2a; border-radius: 4px;
        padding: 0.9rem; font-family: 'IBM Plex Mono', monospace; font-size: 0.72rem;
        color: #4a6a8a; max-height: 260px; overflow-y: auto; line-height: 1.8;
    }
 
    .stButton > button {
        background: #1a3a6e; color: #a8d4ff; border: 1px solid #2a5aaa;
        font-family: 'IBM Plex Mono', monospace; font-size: 0.8rem;
        letter-spacing: 0.05em; border-radius: 4px; transition: background 0.15s;
    }
    .stButton > button:hover { background: #1f4785; border-color: #3a7bd5; }
 
    .stTextInput input, .stSelectbox select, .stTextArea textarea {
        background: #0a0f1a !important; color: #d4dbe8 !important;
        border: 1px solid #1e2d45 !important; font-family: 'IBM Plex Sans', sans-serif;
    }
 
    .meta-tag {
        display: inline-block; background: #0d1825; border: 1px solid #1e3050;
        border-radius: 3px; padding: 1px 8px; font-size: 0.7rem;
        font-family: 'IBM Plex Mono', monospace; color: #5a8ab5; margin-right: 6px;
    }
 
    .tv-link {
        display: inline-block; background: #0d1e35; border: 1px solid #2255aa;
        border-radius: 4px; padding: 4px 10px; font-size: 0.72rem;
        font-family: 'IBM Plex Mono', monospace; color: #7eb8f7;
        text-decoration: none; margin: 3px;
    }
 
    hr { border-color: #141c2a; }
    .stDownloadButton > button { font-size: 0.75rem; padding: 4px 12px; }
    .section-label { color: #4a6a8a; font-size: 0.72rem; font-family: 'IBM Plex Mono', monospace;
                     letter-spacing: 0.08em; margin-bottom: 0.4rem; }
</style>
""", unsafe_allow_html=True)

# ------------------------------------------------------
# Constants
# ------------------------------------------------------

OPENAI_MODEL = "gpt-4.1"
ANTHROPIC_MODEL = "claude-sonnet-4-6"
FRED_BASE = "https://api.stlouisfed.org/fred/series/observations"  ## FRED documentation is - https://fred.stlouisfed.org/docs/api/fred/ 

FX_PAIRS = {
    "USD/GBP": {"symbol": "DEXUSUK", "tv": "FX:GBPUSD", "invent":True},
    "USD/EUR": {"symbol": "DEXUSEU", "tv": "FX:EURUSD", "invent":True},
    "GBP/EUR": {"symbol": "DEXUSUK", "tv": "FX:EURGBP", "invent":False, "derived": True},
}

#FRED series IDs for economic indicators - US, UK, EU
FRED_SERIES = {
    "CPI / Inflation": {
        "US": ("CPIAUCSL",  "US CPI YoY"),
        "UK": ("GBRCPIALLMINMEI",  "UK CPI YoY"),
        "EU": ("CP0000EZ19M086NEST", "EU CPI YoY"),
    },
    "GDP Growth": {
        "US": ("A191RL1Q225SBEA", "US GDP QoQ"),
        "UK": ("UKNGDP", "UK GDP QoQ"),
        "EU": ("EUNGGDP", "EU GDP QoQ"),
    },
    "Interest Rates": {
        "US": ("FEDFUNDS",   "Fed Funds Rate"),
        "UK": ("IUDSOIA",    "BoE Base Rate"),
        "EU": ("ECBDFR",   "ECB Deposit Rate"),
    },
    "Employmnent / NFP": {
        "US":  ("PAYEMS",     "US Nonfarm Payrolls"),
        "UK": ("LRUNTTTTGBM156S","UK Unemployment Rate"),
        "EU":  ("LRHUTTTTEZM156S","EU Unemployment Rate"),
    },
}

LOOKBACK_OPTIONS = { 
    "1 day": 1,
    "7 days": 7,
    "14 days": 14,  
    "30 days": 30,
    "90 days": 90,
    "6 months": 180,
}

REGION_OPTIONS = ["UK / Europe", "US / UK", "US / Europe", "Global (all three)"]

#Signal categories for web-serch layer
Signal_Categories = {
    "Geopolitical Events":        "geopolitical",
    "Central Bank / Rates":       "central_bank",
    "Trade & Sanctions":          "trade_sanctions",
    "Military / Security":        "military_security",
    "Political Stability":        "political_stability",
    "Economic Releases":          "economic_releases",
}

FX_Search_Templates = {
    "geopolitical": [
        '{event} impact USD GBP EUR exchange rate {date_range}',
        '{event} forex currency market reaction {date_range}',
    ],
    "central_bank": [
        'Federal Reserve Bank of England ECB interest rate decision {date_range}',
        'Fed BoE ECB rate impact dollar pound euro forex {date_range}',
        'Fed dot plot BoE MPC ECB forward guidance outlook {date_range}',
    ],
    "trade_sanctions": [
        '{event} trade sanctions tariff dollar pound euro {date_range}',
        '{event} economic impact GBP USD EUR currency {date_range}',
    ],
    "military_security": [
        '{event} military conflict safe haven dollar {date_range}',
        '{event} geopolitical risk currency pound euro dollar {date_range}',
    ],
    "political_stability": [
        '{event} political crisis UK EU US currency {date_range}',
        '{event} government stability forex market reaction {date_range}',
    ],
    "economic_releases": [
        'US UK EU CPI GDP NFP release forex impact {date_range}',
        'economic data surprise dollar pound euro {date_range}',
    ],
}

# Historical FX precedents injected into Claude's system prompt for calibration
FX_HISTORICAL_CONTEXT = """
Key historical geopolitical FX precedents for calibration:
- Black Wednesday (Sep 1992):     GBP collapsed ~15% vs USD/DEM; UK forced out of ERM; major central bank credibility failure
- Asian Financial Crisis (1997):  USD surged vs global currencies; GBP/USD volatile; risk-off flows dominated emerging markets
- Russian debt default (Aug 1998): USD strengthened sharply; EUR (pre-launch proxy) weakened; global flight to safety
- LTCM crisis (Sep 1998):         USD surged initially on liquidity demand; coordinated central bank response stabilized markets
- Euro introduction (Jan 1999):   EUR/USD launched ~1.17; fell to ~0.85 by 2000 (~25% decline); early credibility concerns
- Dot-com bubble peak/crash (2000–2002): USD peaked then weakened significantly; EUR/USD rose from ~0.85 to ~1.10
- 9/11 attacks (Sep 2001):        USD weakened initially on shock; EUR/USD +1–2% short-term; safe-haven flows mixed
- Iraq War onset (Mar 2003):      USD weakened on geopolitical uncertainty; EUR/USD rose above 1.10
- Fed aggressive rate cuts (2001–2003): USD structurally weakened; EUR/USD multi-year uptrend (~0.85 → 1.25)
- Lehman Brothers collapse (Sep 2008): USD surged sharply (global deleveraging); EUR/USD fell ~5% in weeks; GBP/USD dropped below 1.70 → 1.40 in months
- Eurozone sovereign debt crisis peak (2010–2012): EUR/USD fell from ~1.50 to 1.20; GBP/EUR strengthened significantly as Euro confidence collapsed
- ECB "whatever it takes" (Jul 2012): EUR stabilized and rallied; EUR/USD +2–3% in days; major regime shift in Euro confidence
- SNB removes EUR/CHF peg (Jan 2015): EUR shock spillover; EUR/USD dropped ~2% intraday; extreme volatility across EUR crosses
- Brexit referendum (Jun 2016):   GBP/USD fell ~10% in 24h; GBP/EUR fell ~5%
- US Presidential election (Nov 2016): USD surged on growth expectations; EUR/USD fell ~3% in days; GBP/USD initially volatile then weakened
- COVID lockdowns (Mar 2020):     USD surged (safe-haven); EUR/USD fell to 1.07; GBP/USD fell to 1.15
- Fed emergency rate cuts (Mar 2020): USD initially surged then reversed lower; EUR/USD rebounded from ~1.07 to 1.12 rapidly
- Russia-Ukraine invasion (Feb 2022): EUR weakened sharply vs USD and GBP; USD index +3% in 2 weeks
- US CPI upside shock (Jun 2022): USD surged; EUR/USD dropped ~1.5% intraday; GBP/USD fell in tandem on rate repricing
- Fed pivot to aggressive rate hikes (Jun 2022): USD rallied broadly; EUR/USD hit parity (1.00); GBP/USD fell below 1.20
- UK "mini-budget" crisis (Sep 2022): GBP/USD collapsed to ~1.03 (record low); GBP/EUR fell sharply; UK fiscal credibility shock
- US SVB collapse (Mar 2023):     USD weakened; EUR/USD spiked to 1.10; GBP/USD recovered above 1.22
- US debt-ceiling crisis (May 2023): USD volatile; brief risk-off saw GBP and EUR soften
- BoE surprise hike (Aug 2023):   GBP/USD +0.8% intraday; GBP/EUR +0.4%
- ECB rate pause (Oct 2023):      EUR/USD fell 0.6%; GBP/EUR gained 0.3%
Use these precedents to gauge relative magnitude of the current event's likely FX impact.
"""

# ------------------------------------------------------
#FRED Economic Data Fetcher 
# ------------------------------------------------------
def fetch_fred_series(series_id: str, lookback_days: int = 90) -> dict:
    """ 
    Pull the most recent observation(s) for a FRED Series.
    Returns {"latest": value, "prev": value, "date": date_str} on None on failure.
    """
    api_key = os.environ.get("FRED_API_KEY", "")
    if not api_key:
        return None
    
    observation_start = (datetime.now() - timedelta(days=lookback_days + 120)).strftime("%Y-%m-%d")
    params = {
        "series_id":     series_id,
        "api_key":       api_key,
        "file_type":     "json",
        "observation_start": observation_start,
        "sort_order":     "desc",
        "limit":         2,
    }
    try:
        r = requests.get(FRED_BASE, params=params, timeout=8)
        r.raise_for_status()
        obs = r.json().get("observations", [])
        if not obs:
            return None
        latest = obs[0]
        prev = obs[1] if len(obs) > 1 else None
        return{
            "latest": latest["value"] if latest["value"] != "." else "N/A",
            "date": latest["date"],
            "prev": prev["value"] if prev and prev["value"] != "." else "N/A",
        }
    except Exception:
        return None 
    
def fetch_all_economic_data(lookback_days: int, log_fn) -> dict:
    """
    Fetch all four indicator groups (CPI, GDP, Rates, Employment) for US, UK, EU.
    Returns a nested dict: {indicator_label: {country: {latest, prev, date}}}
    """

    results = {}
    for indicator, countries in FRED_SERIES.items():
        results[indicator] = {}
        for country_code, (series_id, label) in countries.items():
            log_fn(f"📊 FRED: fetching {label}…")
            data = fetch_fred_series(series_id, lookback_days)
            results[indicator][country_code] = {
                "label":  label,
                "series": series_id,
                **( data if data else {"latest": "N/A", "prev": "N/A", "date": "N/A"} )
            }
    return results

def format_econ_data_for_prompt(econ_data: dict) -> str:
    """Serialise fetched FRED data into a compact markdown table for Claude."""
    lines = ["## Structured Economic Data (via FRED API)\n"]
    for indicator, countries in econ_data.items():
        lines.append(f"### {indicator}")
        for country_code, d in countries.items():
            prev_str = f" (prev: {d['prev']})" if d['prev'] != "N/A" else ""
            lines.append(f"- **{country_code}** [{d['label']}]: {d['latest']}{prev_str}  _(as of {d['date']})_")
        lines.append("")
    return "\n".join(lines)

# ------------------------------------------------------
#Live FX Rate Display (open.ar-api.com - no key needed)
# ------------------------------------------------------

def fetch_live_fx() -> dict:
    """ Return latest mid-market rates for our three pairs. falls bck to None on error."""
    try:
        r = requests.get("https://open.er-api.com/v6/ltest/USD", timeout=5)
        data = r.json().get("rates", {})
        gbp = data.get("GBP")
        eur = data.get("EUR")
        if not gbp or not eur:
            return {}
        return {
            "USD/GBP": round(gbp, 4),
            "GBP/EUR": round(eur / gbp, 4),
            "USD/EUR":round(eur, 4),
        }
    except Exception:
        return {}
    
# ------------------------------------------------------
# Helpers
# ------------------------------------------------------

def tradingview_url(tv_symbol: str, lookback_days: int) -> str:
    interval = "W" if lookback_days >= 90 else "D"
    return f"https://www.tradingview.com/chart/?symbol={tv_symbol}&interval={interval}"

def _extract_sources_from_response(response) -> list[str]:
    urls: list[str] = []
    for item in response.output:
        if item.type != "message":
            continue
        for part in item.content:
            if not hasattr(part, "annotations"):
                continue
            for annotation in part.annotations:
                if annotation.type == "url_citation" and annotation.url:
                    urls.append(annotation.url)
    return list(dict.fromkeys(urls))

def generate_docx(briefing:str, event: str, lookback:str) -> bytes:
    doc = Document()
    today = datetime.now().strftime("%B %d, %Y")

    for section in doc.sections:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

    title = doc.add_heading("FOREX GEOPOLITICAL SIGNAL INTELLIGENCE BRIEFING", level=0)
    title.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    doc.add_paragraph(f"Event / Topic: {event}")
    doc.add_paragraph(f"FX Pairs: USD/GBP · GBP/EUR · USD/EUR ")
    doc.add_paragraph(f"Lookback: {lookback}")
    doc.add_paragraph(f"Date: {today}")
    doc.add_paragraph("")

    # Disclaimer
    disclaimer = doc.add_paragraph(
        "Disclaimer: This tool and report provides informational and analytical insights only"
        " and does not constitute financial advice or investment recommendations."
    )
    disclaimer.runs[0].font.italic = True
    disclaimer.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph("")


    for line in briefing.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith("---"):
            doc.add_paragraph()
        elif stripped.startswith(("- ", "* ", "• ")) or stripped[:1] in ("🔴", "🟡", "🟢"):
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_run(p, stripped.lstrip("-*• ").strip())
        elif stripped.startswith("**") and stripped.endswith("**"):
            p = doc.add_paragraph()
            p.add_run(stripped.strip("*")).bold = True
        else:
            p = doc.add_paragraph()
            _add_formatted_run(p, stripped)
 
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
 
 
def _add_formatted_run(paragraph, text: str):
    for i, part in enumerate(text.split("**")):
        if part:
            paragraph.add_run(part).bold = (i % 2 == 1)

# ------------------------------------------------------
# Open AI Research Agent
# ------------------------------------------------------

def openai_research(
    event: str,
    lookback_days: int,
    selected_signals: list[str],
    log_fn,
) -> tuple[str, list[str]]:
    """
    Use GPT-4.1 with web_search_preview to gather live geopolitical +
    economic news research for all selected signal categories.
    """
    client = OpenAI(api_key=os.environ["OPENAI_API_KEY"])
 
    end_date   = datetime.now()
    start_date = end_date - timedelta(days=lookback_days)
    date_range = f"{start_date.strftime('%b %d')}–{end_date.strftime('%b %d, %Y')}"
 
    all_findings: list[str] = []
    all_sources:  list[str] = []
 
    for signal_key in selected_signals:
        templates = FX_Search_Templates.get(signal_key, [])
        label = next((k for k, v in Signal_Categories.items() if v == signal_key), signal_key)
 
        queries = [
            t.format(event=event, date_range=date_range)
            for t in templates[:2]
        ]
 
        log_fn(f"🔵 OpenAI web search [{label}]…")
 
        prompt = (
            f"Date range: {date_range}.\n"
            f"Research the following geopolitical event or topic: '{event}'.\n"
            f"Focus: '{label}' — specifically its impact on USD/GBP, GBP/EUR, and USD/EUR exchange rates.\n"
            f"Search angles:\n"
            + "\n".join(f"  - {q}" for q in queries)
            + "\n\nReturn 3–6 concise bullet points of the most relevant, factual findings. "
            "Cite sources inline (outlet name or domain). Be direct and quantify FX moves where possible."
        )
 
        try:
            response = client.responses.create(
                model=OPENAI_MODEL,
                tools=[{"type": "web_search_preview"}],
                input=prompt,
            )
            text = response.output_text or ""
            all_findings.append(f"### {label}\n{text}\n")
            all_sources.extend(_extract_sources_from_response(response))
        except Exception as e:
            log_fn(f"⚠️  OpenAI [{label}] error: {e}")
 
    return "\n\n".join(all_findings), list(dict.fromkeys(all_sources))

# ------------------------------------------------------
# Claude Synthesis Agent
# ------------------------------------------------------

def claude_synthesis(
    event: str,
    lookback_days: int,
    lookback_label: str,
    region: str,
    selected_signals: list[str],
    raw_research: str,
    econ_data_text: str,
    sources: list[str],
    live_rates: dict,
    log_fn,
    result_placeholder,
) -> str:
 
    client = anthropic.Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])
    today  = datetime.now().strftime("%B %d, %Y")
 
    signal_labels = [k for k, v in Signal_Categories.items() if v in selected_signals]
    source_list   = "\n".join(f"- {s}" for s in sources) if sources else "- (sources embedded in research)"
 
    # TradingView links block
    tv_links = "\n".join(
        f"- {pair}: {tradingview_url(meta['tv'], lookback_days)}"
        for pair, meta in FX_PAIRS.items()
    )
 
    # Live rates block
    if live_rates:
        rates_str = " · ".join(f"{pair}: {rate}" for pair, rate in live_rates.items())
    else:
        rates_str = "Live rates unavailable"
 
    system = f"""\
You are a senior FX market intelligence analyst. Today is {today}.
You specialise in identifying how geopolitical events and macroeconomic data drive moves in USD/GBP, GBP/EUR, and USD/EUR.
Write in direct, intelligence-report style: factual, concise, data-grounded, actionable.
Avoid speculation or marketing language. Use Markdown.
 
{FX_HISTORICAL_CONTEXT}
"""
 
    user_msg = f"""\
## FX Intelligence Request
 
**Geopolitical Event / Topic**: {event}
**FX Pairs**: USD/GBP · GBP/EUR · USD/EUR
**Current Live Rates**: {rates_str}
**Lookback Window**: {lookback_label}
**Region Focus**: {region}
**Signal Categories**: {', '.join(signal_labels)}
**Report Date**: {today}
 
---
 
## Web Research (OpenAI GPT-4.1 with live web search)
 
{raw_research}
 
---
 
{econ_data_text}
 
---
 
## Source URLs
{source_list}
 
---
 
## TradingView Chart Links (include these in the Sources section)
{tv_links}
 
---
 
Write a **Forex Geopolitical Signal Intelligence Briefing** using exactly this structure:
 
### Executive Summary
4–5 sentences. What is the single most important thing an FX trader or risk manager needs to know about how **{event}** is affecting USD/GBP, GBP/EUR, and USD/EUR right now?
 
### Geopolitical Signal Findings
For each selected signal category with meaningful findings, write a subsection:
**[Category name]**
- 3–5 bullet points, ordered by significance
- Cite outlet/source inline where possible (e.g. "per Reuters", "via FT")
- Flag each: 🔴 Risk · 🟡 Watch · 🟢 Opportunity
 
### Economic Context
**Separate section** presenting the structured economic data:
- For each indicator (CPI, GDP, Interest Rates, Employment), summarise the latest US / UK / EU readings in one line each
- Note any surprises vs prior readings
- Flag which readings are most FX-relevant right now and why
 
### Cross-Reference: Event × Economics
**The analytical heart of the briefing.**
2–4 bullet points explaining how the geopolitical event interacts with the current economic backdrop:
- Does the economic data amplify or dampen the event's FX impact?
- Example structure: "The [event] shock hits at a moment when [indicator] is [reading], which [amplifies/limits] [pair] moves because [mechanism]."
- Reference historical precedents where relevant (use the calibration examples in your system prompt)
 
### FX Pair Outlook (next 30–90 days)
One short paragraph per pair:
- **USD/GBP**: direction bias, key drivers, level to watch
- **GBP/EUR**: direction bias, key drivers, level to watch
- **USD/EUR**: direction bias, key drivers, level to watch
 
### Recommended Actions
2–4 short, concrete actions for an FX risk manager or trader to consider in the next 30–90 days.
 
### Signals to Monitor
3–4 specific upcoming events, data releases, or central bank decisions worth tracking.
 
### Sources & Charts
List outlets/URLs from grounded sources, then list the three TradingView chart links with a note on the timeframe interval used.
 
---
Total length: 650–900 words. Be concise, direct, and quantitative where possible. Prioritise actionability.
"""
 
    log_fn("🧠 Claude: synthesising FX intelligence briefing…")
 
    full_text = ""
    with client.messages.stream(
        model=ANTHROPIC_MODEL,
        max_tokens=3000,
        system=system,
        messages=[{"role": "user", "content": user_msg}],
    ) as stream:
        for delta in stream.text_stream:
            full_text += delta
            result_placeholder.markdown(
                f'<div class="briefing-wrap">{full_text}</div>',
                unsafe_allow_html=True,
            )
 
    result_placeholder.markdown(
        f'<div class="briefing-wrap">{full_text}</div>',
        unsafe_allow_html=True,
    )
    log_fn("🧠 Claude briefing complete.")
    return full_text
 
# ------------------------------------------------------
# Orchestration 
# ------------------------------------------------------

def run_pipeline(event, lookback_days, lookback_label, region, selected_signals, live_rates, log_fn, result_placement):
 
    log_fn(f"🚀 FX intel request: '{event}'")
    log_fn(f"    Pairs: USD/GBP · GBP/EUR · USD/EUR  |  Window: {lookback_label}")
    log_fn(f"    Signals: {len(selected_signals)} categories active")
 
    # 1. FRED economic data (structured)
    log_fn("📊 Fetching structured economic data from FRED…")
    econ_data = fetch_all_economic_data(lookback_days, log_fn)
    econ_data_text = format_econ_data_for_prompt(econ_data)
    log_fn("✅ FRED data fetched.")
 
    # 2. OpenAI web research
    raw_research, sources = openai_research(event, lookback_days, selected_signals, log_fn)
 
    # 3. Claude synthesis
    briefing = claude_synthesis(
        event, lookback_days, lookback_label, region,
        selected_signals, raw_research, econ_data_text,
        sources, live_rates, log_fn, result_placement,
    )
 
    log_fn(f"✅ Complete — {len(sources)} grounded sources")
    return briefing, sources, econ_data

# ------------------------------------------------------
# UI
# ------------------------------------------------------

st.title("⬡ Forex Geopolitical Signal Intelligence")
st.caption("Cross-reference geopolitical events with economic data to assess FX impact on USD/GBP · GBP/EUR · USD/EUR")

st.markdown(
    '<div style="background:#0d1825; border:1px solid #1e3050; border-radius:4px; '
    'padding:8px 14px; font-size:0.72rem; color:#4a6a8a; '
    'font-family:\'IBM Plex Mono\', monospace; letter-spacing:0.03em; margin-bottom:0.5rem;">'
    '⚠ DISCLAIMER: This tool and report provides informational and analytical insights only '
    'and does not constitute financial advice or investment recommendations.'
    '</div>',
    unsafe_allow_html=True,
)

st.divider()

#------ Live FX Rates---------

live_rates = fetch_live_fx()

if live_rates:
    st.markdown('<div class="section-label">LIVE FX RATES (mid-market)</div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns(3)
    for col, (pair, rate) in zip([c1, c2, c3], live_rates.items()):
        with col:
            st.markdown(
               f'<div class="pair-card">'
                f'<div class="pair-name">{pair}</div>'
                f'<div class="pair-rate">{rate}</div>'
                f'</div>',
                unsafe_allow_html=True,
            )
    st.markdown("") 
            
#------ Input --------------
# 
col1, col2, col3 = st.columns([3, 1, 1])

with col1:
    event = st.text_input(
        "Geopolitical Event / Topic",
         placeholder =  "e.g. Russia-Ukraine conflict, Brexit, Fed rate decision, NATO summit, US-China tariff escalation, UK general election ",
    )

with col2:
    lookback_label = st.selectbox("Lookback Window", list(LOOKBACK_OPTIONS.keys()), index=1)
    lookback_days = LOOKBACK_OPTIONS[lookback_label]

with col3:
    region = st.selectbox("Region Focus", REGION_OPTIONS)

# ------ Signal Category Selection ---------

st.markdown("**Signal Categories**")

if "selected_signals" not in st.session_state:
    st.session_state.selected_signals = list(Signal_Categories.values())

cols = st.columns(len(Signal_Categories))
for i, (label, key) in enumerate(Signal_Categories.items()):
    with cols[i]:
        checked = st.checkbox(label, value=key in st.session_state.selected_signals, key=f"sig_{key}")
        if checked and key not in st.session_state.selected_signals:
            st.session_state.selected_signals.append(key)
        elif not checked and key in st.session_state.selected_signals:
            st.session_state.selected_signals.remove(key)

# ---Economic Indicators Preview------------

with st.expander("📊 Economic Indicators (FRED API) — click to preview available data", expanded=False):
    st.caption("These four indicator groups will be fetched live from FRED and injected into the briefing.")
    for indicator, countries in FRED_SERIES.items():
        st.markdown(f"**{indicator}**")
        for country_code, (series_id, label) in countries.items():
            st.markdown(f"&nbsp;&nbsp;&nbsp;`{country_code}` — {label} &nbsp;·&nbsp; FRED: `{series_id}`")
 
st.divider()

#---TradingView Quick Links------------

st.markdown('<div class="section-label">TRADINGVIEW CHARTS</div>', unsafe_allow_html=True)
tv_cols = st.columns(3)
for col, (pair, meta) in zip(tv_cols, FX_PAIRS.items()):
    with col:
        url = tradingview_url(meta["tv"], lookback_days)
        st.markdown(f'<a class="tv-link" href="{url}" target="_blank">📈 {pair} on TradingView ↗</a>', unsafe_allow_html=True)
 
st.markdown("")
 
run_btn = st.button("⬡ Run FX Geopolitical Scan", type="primary", use_container_width=True)
 

 # ---RUN --------------------------------------------------

if run_btn:
    if not event.strip():
        st.warning("Please enter a geopoliticl event or topics.")
        st.stop()

    if not st.session_state.selected_signals:
        st.warning("Please select at least one signal category.")
        st.stop()

    missing_keys = [k for k in ("ANTHROPIC_API_KEY", "OPENAI_API_KEY") if not os.environ.get(k)]
    if missing_keys:
        st.error(f"Missing environment variables: {', '.join(missing_keys)}")
        st.stop()

    if not os.environ.get("FRED_API_KEY"):
        st.warning("⚠️  FRED_API_KEY not set — economic data will fall back to web search only.")

    st.markdown(
        f'<span class="meta-tag">EVENT</span> **{event}** &nbsp;'
        f'<span class="meta-tag">WINDOW</span> {lookback_label} &nbsp;'
        f'<span class="meta-tag">REGION</span> {region} &nbsp;'
        f'<span class="meta-tag">SIGNALS</span> {len(st.session_state.selected_signals)} active &nbsp;'
        f'<span class="meta-tag">{datetime.now().strftime("%Y-%m-%d %H:%M")}</span>',
        unsafe_allow_html=True,
    )
 
    st.markdown("")
    lcol, rcol = st.columns([1, 2])
 
    with lcol:
        st.markdown("**Pipeline Log**")
        log_box = st.empty()
 
    with rcol:
        st.markdown("**FX Intelligence Briefing**")
        result_box = st.empty()
 
    log_lines = []
 
    def log(msg: str):
        log_lines.append(msg)
        log_box.markdown(
            '<div class="log-box">' + "<br>".join(log_lines[-20:]) + '</div>',
            unsafe_allow_html=True,
        )
 
    try:
        briefing, sources, econ_data = run_pipeline(
            event.strip(), lookback_days, lookback_label, region,
            list(st.session_state.selected_signals),
            live_rates, log, result_box,
        )
 
        st.divider()
 
        # ── Economic Data Summary Table ──────────────────────────────────
        with st.expander("📊 Structured Economic Data (FRED)", expanded=False):
            for indicator, countries in econ_data.items():
                st.markdown(f"**{indicator}**")
                rows = []
                for country_code, d in countries.items():
                    rows.append({
                        "Country": country_code,
                        "Series":  d["label"],
                        "Latest":  d["latest"],
                        "Prev":    d["prev"],
                        "As of":   d["date"],
                    })
                st.table(rows)
            
        # # ── Downloads ───────────────────────────────────────────────────
        base_filename = (
            f"fx_sigint_{event[:30].replace(' ', '_')}"
            f"_{datetime.now().strftime('%Y%m%d_%H%M')}"
        )
 
        dl1, dl2, _ = st.columns([1, 1, 2])
 
        with dl1:
            docx_bytes = generate_docx(briefing, event.strip(), lookback_label)
            st.download_button(
                "⬇️ Download Briefing (.docx)",
                data=docx_bytes,
                file_name=f"{base_filename}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
 
        with dl2:
            st.download_button(
                "⬇️ Download Briefing (.md)",
                data=briefing,
                file_name=f"{base_filename}.md",
                mime="text/markdown",
            )
 
    except Exception as e:
        st.error(f"❌ {type(e).__name__}: {e}")
        st.stop()




    
